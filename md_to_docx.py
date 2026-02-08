
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os

# Mapping: Source Outline -> Revised Content
# Based on the analysis of headers and file content
FILES_MAPPING = [
    {
        "source": "Outline_Chisholm Law _ Supporting Content Rewrite (12.2025).md",
        "revised": "Copy of Chisholm Law _ Supporting Content Rewrite (12.2025)_revised.md",
        "output": "Chisholm_Foundation_vs_Public_Charity.docx"  # H1: Private Foundation vs. Public Charity
    },
    {
        "source": "Outline_Chisholm Law _ Supporting Content Rewrite (12.2025) (1).md",
        "revised": "Copy of Chisholm Law _ Supporting Content Rewrite_revised.md",
        "output": "Chisholm_Nonprofit_101.docx" # H1: Nonprofit 101
    },
    {
        "source": "Outline_Chisholm Law _ Supporting Content Rewrite (12.2025) (2).md",
        "revised": "Copy of Chisholm Law _ Supporting Content_revise.md",
        "output": "Chisholm_Setup_For_Grants.docx" # H1: Setup a Nonprofit to Qualify for Grants
    },
    {
        "source": "Outline_Chisholm Law _ Supporting Content Rewrite (12.2025) (3).md",
        "revised": "Copy of Chisholm Law _ Supporting Content_revisedd.md",
        "output": "Chisholm_Not_For_Profit_Differences.docx" # H1: What Is a Not-For-Profit?
    },
    {
        "source": "Outline_Chisholm Law _ Supporting Content Rewrite (12.2025) (4).md",
        "revised": "Copy of Chisholm Law _ Supporting Content Rewrite_reviseddd.md",
        "output": "Chisholm_Converting_For_Profit.docx" # H1: Converting a For-Profit
    }
]

def extract_metadata(source_file):
    """
    Extracts title tags and meta descriptions from the source markdown file.
    Assumes format:
    **Primary:** Value
    **Secondary:** Value
    """
    with open(source_file, 'r', encoding='utf-8') as f:
        content = f.read()

    metadata = {
        "Primary Title Tag": "",
        "Secondary Title Tag": "",
        "Primary Meta Description": "",
        "Secondary Meta Description": ""
    }

    # Extracting using simple regex looking for the patterns
    # Note: The file structure seems to have sections for "Optimized SEO Title Tag" and "Meta Description"
    
    # Simple state machine or just searching for lines
    lines = content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        if "Optimized SEO Title Tag" in line:
            current_section = "Title"
        elif "Meta Description" in line:
            current_section = "Meta"
        
        if current_section == "Title":
            if "**Primary:**" in line:
                metadata["Primary Title Tag"] = line.replace("**Primary:**", "").strip()
            elif "**Secondary:**" in line:
                metadata["Secondary Title Tag"] = line.replace("**Secondary:**", "").strip()
        elif current_section == "Meta":
            if "**Primary:**" in line:
                metadata["Primary Meta Description"] = line.replace("**Primary:**", "").strip()
            elif "**Secondary:**" in line:
                metadata["Secondary Meta Description"] = line.replace("**Secondary:**", "").strip()
                
    return metadata

def add_markdown_content(doc, content):
    """
    Parses markdown content line by line and adds it to the docx.
    Supports Headers (#) and plain text.
    Bold/Italic formatting is applied via regex.
    """
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        
        # Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=4)
        else:
            # Regular paragraph
            if not stripped:
                continue # Skip empty lines
                
            p = doc.add_paragraph()
            
            # Simple bold properties handling: **text**
            # This is a naive parser. For full markdown support we'd need a real parser.
            # But we want to preserve content verbatim as much as possible.
            
            # To apply bolding, we need to split by **
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Check for links [text](url) - try to keep them readable or just text
                    # keeping it simple: just text
                    p.add_run(part)

def create_docx(mapping):
    source_path = mapping["source"]
    revised_path = mapping["revised"]
    output_filename = mapping["output"]
    
    print(f"Processing {output_filename}...")
    
    # 1. Get Metadata
    metadata = extract_metadata(source_path)
    
    # 2. Get Revised Content
    with open(revised_path, 'r', encoding='utf-8') as f:
        revised_content = f.read()

    # 3. Create Document
    doc = docx.Document()
    
    # 4. Create Metadata Table
    # Table headers: Primary Title Tag, Secondary Title Tag, Primary Meta Description, Secondary Meta Description
    # Re-reading prompt: "table containing Primary Title Tag, Secondary Title Tag, Primary Meta Description, and Secondary Meta Description"
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Primary Title Tag'
    hdr_cells[1].text = 'Secondary Title Tag'
    hdr_cells[2].text = 'Primary Meta Description'
    hdr_cells[3].text = 'Secondary Meta Description'
    
    # Bold headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    row_cells = table.rows[1].cells
    row_cells[0].text = metadata["Primary Title Tag"]
    row_cells[1].text = metadata["Secondary Title Tag"]
    row_cells[2].text = metadata["Primary Meta Description"]
    row_cells[3].text = metadata["Secondary Meta Description"]
    
    doc.add_paragraph() # Spacer
    
    # 5. Add Content
    add_markdown_content(doc, revised_content)
    
    # 6. Save
    doc.save(output_filename)
    print(f"Saved {output_filename}")

def main():
    for mapping in FILES_MAPPING:
        try:
            create_docx(mapping)
        except Exception as e:
            print(f"Error processing {mapping['output']}: {e}")

if __name__ == "__main__":
    main()
